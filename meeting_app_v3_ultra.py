"""
HỆ THỐNG TÓM TẮT BIÊN BẢN CUỘC HỌP - TÍCH HỢP ĐẦY ĐỦ v2
(Tích hợp Silero VAD, xử lý chunk nâng cao, repetition removal từ final_real)
"""

# ========================= CẤU HÌNH =========================
WHISPER_MODEL_PATH   = r"..\whisper_finetuned.pt"
QWEN_MODEL_PATH      = r"..\Qwen_Model_3B"   # thư mục chứa safetensors + config
QWEN_QUANT_CACHE     = "qwen_int8_quantized_cache.pt"  # cache CPU (INT8)
TARGET_CHUNK_SECONDS = 20   # Độ dài mục tiêu mỗi chunk âm thanh (giây)
# ============================================================
FFMPEG_BINARY_PATH  = r"D:\Thanh Va Dat\ffmpeg\bin\ffmpeg.exe"

import customtkinter as ctk
from tkinter import filedialog
import tkinter.messagebox as msgbox
from datetime import datetime
import threading
import traceback
import os
import re
import time
import numpy as np
import torch

# Tự động chọn GPU nếu có, ngược lại dùng CPU
DEVICE = "cuda" 

# if torch.cuda.is_available() else "cpu"

if FFMPEG_BINARY_PATH and os.path.exists(FFMPEG_BINARY_PATH):
    ffmpeg_dir = os.path.dirname(FFMPEG_BINARY_PATH)
    if ffmpeg_dir not in os.environ["PATH"]:
        os.environ["PATH"] = ffmpeg_dir + os.pathsep + os.environ["PATH"]

quant_cache_dir = os.path.dirname(QWEN_QUANT_CACHE)
if quant_cache_dir:
    os.makedirs(quant_cache_dir, exist_ok=True)


_models_loaded  = False
_whisper_model  = None
_qwen_tokenizer = None
_qwen_model     = None
_vad_model      = None          # THÊM MỚI: Silero VAD
_get_speech_ts  = None          # THÊM MỚI: hàm get_speech_timestamps từ VAD


# ========================= NẠP MÔ HÌNH =========================

def load_models(log_callback=None):
    global _models_loaded, _whisper_model, _qwen_tokenizer, _qwen_model
    global _vad_model, _get_speech_ts

    if DEVICE == "cuda":
        # Tăng tốc CUDA: tìm thuật toán conv tối ưu nhất cho phần cứng hiện tại
        torch.backends.cudnn.benchmark = True
    else:
        cpu_count = os.cpu_count() or 4
        torch.set_num_threads(cpu_count)
        torch.set_num_interop_threads(max(1, cpu_count // 2))

    def log(msg):
        print(msg)
        if log_callback:
            log_callback(msg)

    # ── Helper: torch.load với memory-mapping nếu PyTorch ≥ 2.1 ─────────────
    def _torch_load(path, map_location):
        """Dùng mmap=True (PyTorch ≥ 2.1) để đọc file trực tiếp từ ổ đĩa
        mà không cần copy toàn bộ vào RAM trước — giảm đáng kể thời gian nạp."""
        ver = tuple(int(x) for x in torch.__version__.split(".")[:2])
        if ver >= (2, 1):
            return torch.load(path, map_location=map_location,
                              weights_only=True, mmap=True)
        return torch.load(path, map_location=map_location, weights_only=True)

    whisper_error = [None]
    qwen_error    = [None]
    vad_error     = [None]

    # ── Whisper ──────────────────────────────────────────────────────────────
    def _load_whisper():
        try:
            import whisper
            log(f"[Whisper] Đang nạp ({DEVICE.upper()})...")
            model = whisper.load_model("small", device=DEVICE)
            checkpoint = _torch_load(WHISPER_MODEL_PATH, map_location=DEVICE)
            if isinstance(checkpoint, dict):
                state_dict = (checkpoint.get("model_state_dict")
                              or checkpoint.get("state_dict")
                              or checkpoint)
            else:
                state_dict = checkpoint
            if any(k.startswith("model.") for k in list(state_dict.keys())[:10]):
                state_dict = {k.replace("model.", ""): v for k, v in state_dict.items()}
            model.load_state_dict(state_dict, strict=False)
            model.eval()
            global _whisper_model
            _whisper_model = model
            log(f"[Whisper] ✔ Sẵn sàng")
        except Exception as e:
            whisper_error[0] = str(e)
            log(f"[Whisper] ✘ Lỗi: {e}")

    # ── Qwen ─────────────────────────────────────────────────────────────────
    def _load_qwen():
        """
        Nạp Qwen từ thư mục safetensors chuẩn HuggingFace.
        - GPU  : load FP16 trực tiếp từ QWEN_MODEL_PATH (safetensors, không cần cache riêng)
        - CPU  : lần đầu quantize INT8 và lưu cache .pt; lần sau load cache để tiết kiệm thời gian
        """
        try:
            from transformers import AutoTokenizer, AutoConfig, AutoModelForCausalLM
            global _qwen_tokenizer, _qwen_model

            log(f"[Qwen] Đang nạp ({DEVICE.upper()})...")
            tokenizer = AutoTokenizer.from_pretrained(
                QWEN_MODEL_PATH, trust_remote_code=True)

            if DEVICE == "cuda":
                model = AutoModelForCausalLM.from_pretrained(
                    QWEN_MODEL_PATH,
                    torch_dtype=torch.float16,
                    device_map="cuda",
                    low_cpu_mem_usage=True,
                    trust_remote_code=True,
                    attn_implementation="sdpa"  # <--- BẬT KÍCH HOẠT TỐI ƯU CỦA ĐỂ RTX 3050 KHÔNG BỊ TRÀN CHAI CỔ
                )
            else:
                # CPU: dung INT8 quantization de giam RAM
                if os.path.exists(QWEN_QUANT_CACHE):
                    config = AutoConfig.from_pretrained(
                        QWEN_MODEL_PATH, trust_remote_code=True)
                    model = AutoModelForCausalLM.from_config(
                        config, trust_remote_code=True)
                    model = torch.quantization.quantize_dynamic(
                        model, {torch.nn.Linear}, dtype=torch.qint8)
                    state_dict = _torch_load(QWEN_QUANT_CACHE, map_location="cpu")
                    model.load_state_dict(state_dict, strict=False)
                else:
                    log("[Qwen] Lần đầu chạy — đang quantize INT8, có thể mất vài phút...")
                    model = AutoModelForCausalLM.from_pretrained(
                        QWEN_MODEL_PATH,
                        torch_dtype=torch.float32,
                        low_cpu_mem_usage=True,
                        trust_remote_code=True
                    )
                    model = torch.quantization.quantize_dynamic(
                        model, {torch.nn.Linear}, dtype=torch.qint8)
                    torch.save(model.state_dict(), QWEN_QUANT_CACHE)

            model.eval()
            _qwen_tokenizer = tokenizer
            _qwen_model     = model
            # torch.compile tăng tốc inference ~20-40% trên GPU (PyTorch 2.0+)
            # mode='reduce-overhead': tối ưu cho inference lặp lại nhiều lần
            ver = tuple(int(x) for x in torch.__version__.split(".")[:2])
            if ver >= (2, 0) and DEVICE == "cuda":
                try:
                    model = torch.compile(model, mode="reduce-overhead")
                    log("[Qwen] torch.compile() kích hoạt")
                except Exception:
                    pass  # Bỏ qua nếu compile không hỗ trợ

            log(f"[Qwen] ✔ Sẵn sàng")
        except Exception as e:
            qwen_error[0] = str(e)
            log(f"[Qwen] ✘ Lỗi: {e}")

    # ── Silero VAD ────────────────────────────────────────────────────────────
    def _load_vad():
        try:
            global _vad_model, _get_speech_ts
            log("[VAD] Đang nạp...")
            vad_m, utils = torch.hub.load(
                repo_or_dir="snakers4/silero-vad",
                model="silero_vad",
                trust_repo=True,
                force_reload=False
            )
            if DEVICE == "cuda":
                vad_m = vad_m.to("cuda")
            _vad_model     = vad_m
            _get_speech_ts = utils[0]
            log("[VAD] ✔ Sẵn sàng")
        except Exception as e:
            vad_error[0] = str(e)
            log(f"[VAD] ✘ Không nạp được, bỏ qua lọc VAD.")

    log(f"Đang nạp mô hình AI (thiết bị: {DEVICE.upper()})...")
    t1 = threading.Thread(target=_load_whisper, daemon=True)
    t2 = threading.Thread(target=_load_qwen,    daemon=True)
    t3 = threading.Thread(target=_load_vad,     daemon=True)
    t1.start(); t2.start(); t3.start()
    t1.join();  t2.join();  t3.join()

    if DEVICE == "cuda":
        torch.cuda.empty_cache()

    if whisper_error[0] or qwen_error[0]:
        return False

    _models_loaded = True
    log("✔ Tất cả mô hình đã sẵn sàng!")
    return True


# ====================== FORM BIÊN BẢN CHÍNH THỨC ======================
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt, Inches

def tao_form_bien_ban_mau(full_text):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # Quét từng dòng văn bản mà người dùng đã chỉnh sửa trên giao diện
    lines = full_text.split('\n')
    
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            doc.add_paragraph()
            continue
            
        if line_strip == "CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM" or line_strip == "Độc Lập - Tự Do - Hạnh Phúc":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line_strip).bold = True
            
        elif line_strip == "BIÊN BẢN CUỘC HỌP":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line_strip)
            run.bold = True
            run.font.size = Pt(16)
            p.paragraph_format.space_after = Pt(24)
            
        elif line_strip.startswith("Thời gian tiến hành:") or line_strip.startswith("Thành phần tham gia:") or line_strip.startswith("Nội dung cuộc họp:"):
            p = doc.add_paragraph()
            idx = line.find(":")
            if idx != -1:
                p.add_run(line[:idx+1]).bold = True
                p.add_run(line[idx+1:])
            else:
                p.add_run(line).bold = True
            p.paragraph_format.space_after = Pt(6)
            if line_strip.startswith("Nội dung cuộc họp:"):
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.keep_with_next = True 

        elif "Người lập biên bản" in line_strip and "Chủ toạ" in line_strip:
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].allow_break_across_pages = False
            
            cell_left = table.cell(0, 0)
            p_left1 = cell_left.paragraphs[0]
            p_left1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_left1.add_run("Người lập biên bản").bold = True
            p_left2 = cell_left.add_paragraph()
            p_left2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_left2.add_run("(Ký, ghi rõ họ tên)").italic = True

            cell_right = table.cell(0, 1)
            p_right1 = cell_right.paragraphs[0]
            p_right1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_right1.add_run("Chủ tọa").bold = True
            p_right2 = cell_right.add_paragraph()
            p_right2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_right2.add_run("(Ký, ghi rõ họ tên)").italic = True

        elif line_strip.startswith("Cuộc họp kết thúc vào lúc"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.add_run(line_strip)
            p.paragraph_format.keep_with_next = True
            
        elif (line_strip.startswith('[') and line_strip.endswith(']:')) or line_strip.startswith('📌'):
            p = doc.add_paragraph()
            run = p.add_run(line_strip)
            run.bold = True
            p.paragraph_format.space_before = Pt(12) 
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.keep_with_next = True 
            
        else:
            p = doc.add_paragraph()
            p.add_run(line_strip)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
            if line_strip.startswith("-"):
                 p.paragraph_format.left_indent = Inches(0.25)
            else:
                 p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.2 
            p.paragraph_format.keep_together = True 

    return doc


# ====================== XỬ LÝ ÂM THANH (NÂNG CẤP TỪ NOTEBOOK) ======================

def reduce_audio_noise(audio_input, sr=None):
    """Giảm nhiễu nền cho file âm thanh."""
    import librosa
    import noisereduce as nr
    if isinstance(audio_input, str):
        audio_data, sampling_rate = librosa.load(audio_input, sr=sr)
    else:
        audio_data, sampling_rate = audio_input, sr
    return nr.reduce_noise(y=audio_data, sr=sampling_rate), sampling_rate


def split_audio_into_chunks(reduced_audio, sr, target_chunk_seconds=TARGET_CHUNK_SECONDS):
    """
    Chia âm thanh thành các chunk dựa trên khoảng lặng.
    (Nâng cấp từ notebook: dùng detect_silence + target_chunk_seconds thay vì split_on_silence)
    """
    from pydub import AudioSegment
    from pydub.silence import detect_silence

    if FFMPEG_BINARY_PATH:
        AudioSegment.converter = FFMPEG_BINARY_PATH


    audio_int16 = (reduced_audio * 32767).astype(np.int16)
    audio_segment = AudioSegment(
        audio_int16.tobytes(),
        frame_rate=sr,
        sample_width=2,
        channels=1
    )

    silences = detect_silence(
        audio_segment,
        min_silence_len=300,
        silence_thresh=-40,
        seek_step=10
    )

    chunks = []
    chunk_start = 0
    target_ms  = target_chunk_seconds * 1000
    min_chunk_ms = 5000
    max_chunk_ms = 30000

    for silence_start, silence_end in silences:
        silence_mid   = (silence_start + silence_end) // 2
        chunk_duration = silence_mid - chunk_start

        if chunk_duration >= target_ms:
            if chunk_duration > max_chunk_ms:
                while chunk_start < silence_mid:
                    end = min(chunk_start + max_chunk_ms, silence_mid)
                    chunks.append(audio_segment[chunk_start:end])
                    chunk_start = end
            else:
                chunks.append(audio_segment[chunk_start:silence_mid])
                chunk_start = silence_mid

    if chunk_start < len(audio_segment):
        last_chunk = audio_segment[chunk_start:]
        if len(last_chunk) >= min_chunk_ms:
            chunks.append(last_chunk)
        elif chunks:
            chunks[-1] = chunks[-1] + last_chunk
        else:
            chunks.append(last_chunk)

    return chunks


# ── THÊM MỚI: Lọc khoảng lặng bằng Silero VAD ──────────────────────────────
def remove_silence_vad(audio_array, sample_rate=16000):
    """
    Dùng Silero VAD để loại bỏ các đoạn im lặng bên trong chunk.
    Nếu VAD chưa tải được, trả về nguyên mảng gốc.
    """
    if _vad_model is None or _get_speech_ts is None:
        return audio_array  # VAD không có → bỏ qua

    audio_tensor = torch.from_numpy(audio_array)
    # Đưa tensor lên cùng device với VAD model (CPU hoặc CUDA)
    vad_device = next(_vad_model.parameters()).device
    audio_tensor = audio_tensor.to(vad_device)

    speech_timestamps = _get_speech_ts(
        audio_tensor,
        _vad_model,
        sampling_rate=sample_rate
    )

    if len(speech_timestamps) == 0:
        return np.array([], dtype=np.float32)

    # Timestamp trả về là index trên CPU numpy array gốc — dùng trực tiếp
    speech_audio = [audio_array[ts["start"]:ts["end"]] for ts in speech_timestamps]
    return np.concatenate(speech_audio)


# ── THÊM MỚI: Bộ hàm loại bỏ lặp từ/câu (repetition removal) ───────────────
def remove_context_overlap(text, previous_context, min_overlap=3):
    """Loại bỏ phần đầu của text trùng với cuối previous_context."""
    text_words    = text.split()
    context_words = previous_context.split()

    if len(context_words) < min_overlap or len(text_words) < min_overlap:
        return text

    max_overlap = min(len(context_words), len(text_words), 20)
    for overlap_len in range(max_overlap, min_overlap - 1, -1):
        if context_words[-overlap_len:] == text_words[:overlap_len]:
            return " ".join(text_words[overlap_len:])
    return text


def remove_immediate_word_repetition(text, max_consecutive=2):
    """Loại bỏ các từ lặp liên tiếp (ví dụ: 'và và và' → 'và và')."""
    words = text.split()
    if not words:
        return text

    result = [words[0]]
    consecutive_count = 1
    for word in words[1:]:
        if word == result[-1]:
            consecutive_count += 1
            if consecutive_count <= max_consecutive:
                result.append(word)
        else:
            result.append(word)
            consecutive_count = 1
    return " ".join(result)


def remove_phrase_repetition_advanced(text, min_phrase_len=2, max_phrase_len=10):
    """Loại bỏ cụm từ lặp liên tiếp (ví dụ: 'hôm nay hôm nay' → 'hôm nay')."""
    words = text.split()
    n = len(words)
    if n < min_phrase_len * 2:
        return text

    i = 0
    result = []
    while i < n:
        found = False
        for phrase_len in range(min(max_phrase_len, n - i), min_phrase_len - 1, -1):
            if i + phrase_len * 2 > n:
                continue
            phrase1 = words[i: i + phrase_len]
            phrase2 = words[i + phrase_len: i + phrase_len * 2]
            if phrase1 == phrase2:
                result.extend(phrase1)
                j = i + phrase_len * 2
                while j + phrase_len <= n and words[j: j + phrase_len] == phrase1:
                    j += phrase_len
                i = j
                found = True
                break
        if not found:
            result.append(words[i])
            i += 1
    return " ".join(result)


def remove_fillers(text):
    """Loại bỏ các từ đệm không có nghĩa."""
    fillers = {"ờ", "à", "ừ", "ừm", "uh", "ờm", "ơ", "um"}
    return " ".join(w for w in text.split() if w.lower() not in fillers)


def remove_sentence_repetition(text):
    """Loại bỏ các câu bị lặp lại."""
    sentences = text.split(". ")
    cleaned, seen = [], set()
    for s in sentences:
        s_norm = s.strip().lower()
        if s_norm not in seen:
            cleaned.append(s)
            seen.add(s_norm)
    return ". ".join(cleaned)


def clean_text_with_validation(text):
    """Áp dụng toàn bộ bộ lọc lặp trước khi đưa vào Qwen."""
    text = remove_immediate_word_repetition(text, max_consecutive=2)
    text = remove_phrase_repetition_advanced(text)
    text = remove_sentence_repetition(text)
    text = remove_fillers(text)
    return text


def detect_and_fix_repetition_realtime(text, previous_context=""):
    """Gọi trong vòng lặp xử lý chunk để fix lặp real-time."""
    if previous_context:
        text = remove_context_overlap(text, previous_context)
    text = remove_immediate_word_repetition(text)
    text = remove_phrase_repetition_advanced(text)
    return text.strip()


# ── process_audio_chunks – NÂNG CẤP từ notebook ─────────────────────────────
def process_audio_chunks(chunks, whisper_model, progress_callback=None):
    """
    Xử lý từng chunk âm thanh với Whisper.
    Nâng cấp so với phiên bản cũ:
      - Dùng Silero VAD lọc im lặng trong mỗi chunk
      - Gộp các chunk ngắn < 5s
      - Dùng previous_context để Whisper có ngữ cảnh liên tục
      - Nhiều temperature để Whisper tự chọn kết quả tốt nhất
      - Fix lặp realtime sau mỗi chunk
    """
    processed_texts  = []
    total_chunks     = len(chunks)
    previous_context = ""
    max_context_words = 50
    current_audio_time_ms = 0
    pending_samples  = np.array([], dtype=np.float32)

    use_fp16 = (DEVICE == "cuda")

    for i, chunk in enumerate(chunks):
        chunk_duration_ms = len(chunk)
        start_sec = current_audio_time_ms / 1000.0
        end_sec   = (current_audio_time_ms + chunk_duration_ms) / 1000.0

        # Chuyển sang float32 16kHz
        chunk_16k = chunk.set_frame_rate(16000).set_channels(1)
        samples   = np.array(chunk_16k.get_array_of_samples(), dtype=np.float32)
        max_val   = float(1 << (8 * chunk_16k.sample_width - 1))
        samples  /= max_val

        # Lọc im lặng bằng VAD
        cleaned_samples = remove_silence_vad(samples)

        if cleaned_samples.size == 0:
            current_audio_time_ms += chunk_duration_ms
            if progress_callback:
                progress_callback(i + 1, total_chunks)
            continue

        pending_samples = np.concatenate([pending_samples, cleaned_samples])
        current_pending_sec = len(pending_samples) / 16000.0

        # Gộp với chunk tiếp theo nếu còn ngắn (< 5s) và chưa phải chunk cuối
        if i < total_chunks - 1 and current_pending_sec < 5.0:
            current_audio_time_ms += chunk_duration_ms
            continue

        with torch.no_grad():
            result = whisper_model.transcribe(
                pending_samples,
                language="vi",
                temperature=[0.0, 0.2, 0.4, 0.6, 0.8],
                beam_size=5,
                best_of=5,
                condition_on_previous_text=True,
                prompt=previous_context if previous_context else None,
                fp16=use_fp16,
                suppress_blank=True
            )

        raw_text     = result["text"].strip()
        cleaned_text = detect_and_fix_repetition_realtime(raw_text, previous_context)
        processed_texts.append(cleaned_text)

        # Cập nhật context
        words = cleaned_text.split()
        previous_context = " ".join(words[-max_context_words:]) if len(words) > max_context_words else cleaned_text

        print(f"  Nhận dạng: {i+1}/{total_chunks} đoạn ({start_sec:.0f}s → {end_sec:.0f}s)")

        pending_samples = np.array([], dtype=np.float32)
        current_audio_time_ms += chunk_duration_ms

        if progress_callback:
            progress_callback(i + 1, total_chunks)

    return processed_texts


def merge_text(texts):
    return " ".join(t for t in texts if t.strip()).strip()


def transcribe_audio_pipeline(audio_path, progress_callback=None):
    """Pipeline hoàn chỉnh: giảm nhiễu → chia chunk → VAD → Whisper → gộp văn bản."""
    print("[Giai đoạn 1/3] Đang nhận dạng giọng nói (Whisper)...")
    reduced, sr = reduce_audio_noise(audio_path)
    chunks = split_audio_into_chunks(reduced, sr, TARGET_CHUNK_SECONDS)
    print(f"  Đã chia thành {len(chunks)} đoạn âm thanh")
    texts  = process_audio_chunks(chunks, _whisper_model, progress_callback)
    return merge_text(texts)


# ====================== LÀM SẠCH VĂN BẢN (NÂNG CẤP TỪ NOTEBOOK) ======================

# System prompt nâng cấp từ notebook
_CLEAN_SYSTEM_PROMPT = """
Bạn là một công cụ tự động hiệu đính văn bản (Text Corrector). 
Nhiệm vụ DUY NHẤT của bạn là nhận văn bản ASR thô, sửa lỗi chính tả và điền thêm 
dấu câu (chấm, phẩy, hỏi chấm) để câu văn đúng ngữ pháp tiếng Việt.

QUY TẮC TUYỆT ĐỐI (FATAL RULES):
GIỮ NGUYÊN 100% ngữ nghĩa và thông tin.
Chỉ sửa từ sai chính tả, tuyệt đối không tự ý thêm bớt hay thay đổi từ vựng của người nói.
KHÔNG trò chuyện, KHÔNG giải thích, KHÔNG chào hỏi.
TUYỆT ĐỐI KHÔNG sử dụng các câu mào đầu như: "Văn bản sau đây...", "Đây là kết quả...", "Dấu chấm:".
ĐẦU RA (OUTPUT) CHỈ ĐƯỢC CHỨA VĂN BẢN ĐÃ SỬA, không chứa bất kỳ ký tự nào khác.

Dưới đây là các ví dụ về cách bạn phải hoạt động:

[Ví dụ 1]
Input: hôm nay thời tiết rất đệp chúng ta có nên đi bơi không nhỉ tớ nghĩ là có
Output: Hôm nay thời tiết rất đẹp, chúng ta có nên đi bơi không nhỉ? Tớ nghĩ là có.

[Ví dụ 2]
Input: chào các bạn hôm nay tôi xẽ trình bày về dự án mới dự án này gổm ba phần phần một là
Output: Chào các bạn. Hôm nay tôi sẽ trình bày về dự án mới. Dự án này gồm ba phần, phần một là...
""".strip()


def _build_clean_prompt(transcript):
    return f"""
Văn bản cuộc họp:

{transcript}

Hãy chỉnh sửa theo đúng các quy tắc đã nêu.
""".strip()


def remove_non_vietnamese(text):
    """Loại bỏ ký tự không phải tiếng Việt."""
    return re.sub(r"[^\u00C0-\u1EF9a-zA-Z0-9\s.,:;!?%]", "", text)


def _clean_chunk_with_qwen_batch(texts):
    """Dùng Qwen sửa chính tả & dấu câu cho MỘT LÔ (Batch) các đoạn văn bản."""
    # Bắt buộc đệm (padding) bên trái để GPU tiếp tục sinh chữ ở bên phải
    _qwen_tokenizer.padding_side = "left"
    if _qwen_tokenizer.pad_token_id is None:
        _qwen_tokenizer.pad_token = _qwen_tokenizer.eos_token
        
    prompts = []
    for text in texts:
        messages = [
            {"role": "system", "content": _CLEAN_SYSTEM_PROMPT},
            {"role": "user",   "content": _build_clean_prompt(text)}
        ]
        prompt = _qwen_tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
        prompts.append(prompt)

    inputs = _qwen_tokenizer(
        prompts, 
        return_tensors="pt", 
        padding=True, 
        truncation=True, 
        max_length=2500
    ).to(_qwen_model.device)

    with torch.no_grad():
        outputs = _qwen_model.generate(
            **inputs,
            max_new_tokens=200,
            do_sample=False,
            repetition_penalty=1.15,
            eos_token_id=_qwen_tokenizer.eos_token_id,
            pad_token_id=_qwen_tokenizer.eos_token_id
        )

    results = []
    for i in range(len(prompts)):
        generated_ids = outputs[i][inputs.input_ids.shape[-1]:]
        response = _qwen_tokenizer.decode(generated_ids, skip_special_tokens=True)
        results.append(response.strip())
    return results


def restore_punctuation_and_clean(text):
    """
    Làm sạch và khôi phục dấu câu cho văn bản ASR thô.
    """
    # Bước 1: loại bỏ lặp từ/câu
    text = clean_text_with_validation(text)

    # Bước 2: chia thành chunk 120 từ
    words  = text.split()
    chunks = [" ".join(words[i:i + 120]) for i in range(0, len(words), 120)]
    total  = len(chunks)
    print(f"  Làm sạch văn bản: {total} đoạn cần xử lý")

    # BATCH SIZE CHO RTX 3050: Gộp 4 đoạn gọi GPU 1 lần
    BATCH_SIZE = 4
    cleaned_chunks = []
    
    # Lọc bỏ các chunk quá rác trước khi đẩy vào batch
    valid_chunks = [c for c in chunks if len(c.strip()) >= 5]
    total_valid = len(valid_chunks)

    for i in range(0, total_valid, BATCH_SIZE):
        batch = valid_chunks[i : i + BATCH_SIZE]
        print(f"  Làm sạch Batch {i//BATCH_SIZE + 1} (gồm {len(batch)} đoạn)...")
        
        batch_cleaned = _clean_chunk_with_qwen_batch(batch)
        
        for cleaned in batch_cleaned:
            cleaned = remove_non_vietnamese(cleaned)
            cleaned_chunks.append(cleaned)

    return "\n".join(cleaned_chunks)


# ====================== TÓM TẮT (giữ nguyên) ======================

def generate_summary_batch(texts, max_new_tokens=200, is_final=False):
    """
    Kích hoạt khả năng đẩy hàng loạt câu vào Qwen cùng một thời điểm.
    Tối đa hoá hàng ngàn nhân CUDA của card rời thay vì bắt nó chờ.
    """
    if is_final:
        system_prompt = """Bạn là Thư ký cuộc họp. Hãy tóm tắt văn bản thành danh sách gạch đầu dòng phẳng.
<rules>
1. Bắt đầu mỗi ý bằng ký tự "-". KHÔNG dán nhãn chủ đề.
2. KHÔNG ghi các từ mào đầu như: "Vấn đề:", "Ý kiến:", "Kết luận:".
3. TRÌNH TỰ THỜI GIAN: Tóm tắt bao quát toàn bộ sự kiện theo đúng thứ tự.
4. TRÍCH XUẤT TRỰC TIẾP.
</rules>"""
        rep_penalty = 1.05
    else:
        system_prompt = """Trích xuất ý chính từ đoạn hội thoại.
<rules>
- Viết dưới dạng gạch đầu dòng (-).
- LỌC RÁC CỰC MẠNH.
- Nếu toàn rác -> PASS
</rules>"""
        rep_penalty = 1.05

    _qwen_tokenizer.padding_side = "left"
    if _qwen_tokenizer.pad_token_id is None:
        _qwen_tokenizer.pad_token = _qwen_tokenizer.eos_token

    prompts = []
    for text in texts:
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"<van_ban>\n{text}\n</van_ban>\n\nHãy tóm tắt:"}
        ]
        prompt = _qwen_tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
        prompts.append(prompt)

    inputs = _qwen_tokenizer(
        prompts,
        return_tensors="pt",
        padding=True,
        truncation=True,
        max_length=2500
    ).to(_qwen_model.device)

    with torch.no_grad():
        outputs = _qwen_model.generate(
            **inputs,
            max_new_tokens=max_new_tokens if is_final else 200,
            repetition_penalty=rep_penalty,
            do_sample=False,
            pad_token_id=_qwen_tokenizer.eos_token_id,
            eos_token_id=_qwen_tokenizer.eos_token_id
        )

    results = []
    for i in range(len(prompts)):
        generated_ids = outputs[i][inputs.input_ids.shape[-1]:]
        summary = _qwen_tokenizer.decode(generated_ids, skip_special_tokens=True)
        summary = summary.replace("<|im_end|>", "").strip()
        results.append(summary)

    return results

def generate_summary(text, max_new_tokens=1000, is_final=True):
    from underthesea import sent_tokenize
    if is_final:
        system_prompt = """Bạn là Thư ký cuộc họp. Hãy tóm tắt văn bản thành danh sách gạch đầu dòng phẳng.
<rules>
1. Bắt đầu mỗi ý bằng ký tự "-". KHÔNG dán nhãn chủ đề.
2. KHÔNG ghi các từ mào đầu như: "Vấn đề:", "Ý kiến:", "Kết luận:".
3. TRÌNH TỰ THỜI GIAN: Tóm tắt bao quát toàn bộ sự kiện theo đúng thứ tự.
4. TRÍCH XUẤT TRỰC TIẾP: NẾU văn bản đã ngắn gọn và súc tích, hãy trực tiếp trích xuất các ý chính mà không cần thay đổi quá nhiều từ ngữ.
</rules>"""
        rep_penalty = 1.05
    else:
        system_prompt = """Trích xuất ý chính từ đoạn hội thoại.
<rules>
- Viết dưới dạng gạch đầu dòng (-).
- LỌC RÁC CỰC MẠNH (MỆNH LỆNH): Tuyệt đối KHÔNG trích xuất các câu cảm thán, từ đệm (ví dụ: "Dạ vâng", "Chúc mừng").
- NẾU đoạn văn chỉ toàn rác giao tiếp -> TRẢ VỀ ĐÚNG 1 CHỮ: PASS
</rules>"""
        rep_penalty = 1.05

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"<van_ban>\n{text}\n</van_ban>\n\nHãy tóm tắt:"}
    ]

    text_input = _qwen_tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
    inputs = _qwen_tokenizer(text_input, return_tensors="pt", truncation=True, max_length=2500).to(_qwen_model.device)

    with torch.no_grad():
        outputs = _qwen_model.generate(
            **inputs,
            max_new_tokens=max_new_tokens if is_final else 200,
            repetition_penalty=rep_penalty,
            do_sample=False,
            pad_token_id=_qwen_tokenizer.eos_token_id,
            eos_token_id=_qwen_tokenizer.eos_token_id
        )

    generated_ids = [output_ids[len(input_ids):] for input_ids, output_ids in zip(inputs.input_ids, outputs)]
    summary = _qwen_tokenizer.batch_decode(generated_ids, skip_special_tokens=True)[0]

    summary = summary.replace("<|im_end|>", "").strip()

    if not is_final:
        return summary

    # ==========================================
    # XỬ LÝ HẬU KỲ & FALLBACK BẢO VỆ
    # ==========================================
    cleaned_lines = []
    for line in summary.split('\n'):
        line = line.strip()

        if not line or re.search(r'(chào các bạn|chào mọi người|cảm ơn|hẹn gặp lại)', line, re.IGNORECASE):
            continue

        line = re.sub(r"^(- )?(### Biên bản|Biên bản họp|Tóm tắt|Nội dung chính|Vấn đề).*?:?", "", line, flags=re.IGNORECASE).strip()

        # Chặn triệt để lỗi sinh ra dòng chỉ có dấu "-" hoặc dòng rỗng sau khi bị cạo
        if line == "-" or line == "":
            continue

        if line:
            if not line.startswith('-'):
                line = '- ' + line
            cleaned_lines.append(line)

    final_summary = '\n'.join(cleaned_lines).strip()

    # FALLBACK CỨU HỘ: Nếu AI thực sự bị "tịt ngòi" và trả về rỗng, fallback về việc chia câu cơ bản
    if not final_summary:
        print("[Warning] Model bị ngắt kết quả, kích hoạt Fallback chia câu.")
        sentences = sent_tokenize(text)
        return "\n".join(["- " + s for s in sentences])

    return final_summary


def enhanced_textrank(text):
    from underthesea import sent_tokenize, word_tokenize
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import networkx as nx

    sentences = sent_tokenize(text)
    if len(sentences) < 4:
        return generate_summary(text, is_final=True)

    sentences_seg = [word_tokenize(s, format="text") for s in sentences]
    try:
        tfidf = TfidfVectorizer().fit_transform(sentences_seg)
        sim_matrix = cosine_similarity(tfidf)
        graph = nx.from_numpy_array(sim_matrix)
        scores = nx.pagerank(graph)

        ranked = sorted(((scores[i], s) for i, s in enumerate(sentences)), reverse=True)
        top_n = max(4, int(len(sentences) * 0.45))

        # Dùng set để lọc trùng lặp khi thêm câu
        selected_sentences_set = set([s for _, s in ranked[:top_n]])

        # GẮN ĐINH NGỮ CẢNH: Bắt buộc giữ 2 câu đầu tiên bất chấp điểm số TextRank
        for initial_sentence in sentences[:2]:
            selected_sentences_set.add(initial_sentence)

        # Sắp xếp lại danh sách các câu đã chọn theo đúng thứ tự xuất hiện ban đầu
        selected_sentences = sorted(list(selected_sentences_set), key=lambda x: sentences.index(x))
        filtered_text = " ".join(selected_sentences)

        # --- PHẦN CHỈNH SỬA TẠI ĐÂY ---
        # Bước 1: Sinh gạch đầu dòng từ văn bản đã lọc
        flat_summary = generate_summary(filtered_text, is_final=True)

        # Bước 2: Gom nhóm và chuyển đổi thành đoạn văn xuôi
        final_grouped_paragraphs = dynamic_clustering_to_paragraphs(flat_summary)

        return final_grouped_paragraphs
        # ------------------------------

    except Exception as e:
        print(f"Lỗi TextRank: {e}")
        # Fallback trong trường hợp lỗi vẫn đảm bảo ra định dạng đoạn văn
        flat_summary = generate_summary(text, is_final=True)
        return dynamic_clustering_to_paragraphs(flat_summary)


def chunking_method(text):
    from underthesea import sent_tokenize

    sentences = sent_tokenize(text)
    if len(sentences) < 6:
        return generate_summary(text, is_final=True)

    chunk_size = 10  # Mức lý tưởng cho 3B
    stride = 7       # Bước nhảy: Cửa sổ trượt gối đầu 3 câu (Overlap)

    chunks = []
    for i in range(0, len(sentences), stride):
        chunk = " ".join(sentences[i : i + chunk_size])
        chunks.append(chunk)
        # Dừng lại nếu đã chạm đến cuối văn bản
        if i + chunk_size >= len(sentences):
            break

    # Để an toàn cho 4GB RAM, ta chia chunks ra làm nhiều mini-batches
    BATCH_SIZE = 4
    partials = []
    
    print(f"  Tiến hành Tóm tắt {len(chunks)} đoạn qua {len(chunks)//BATCH_SIZE + 1} Batch...")
    
    for i in range(0, len(chunks), BATCH_SIZE):
        batch_chunks = chunks[i : i + BATCH_SIZE]
        print(f"  Đang chạy Batch {i//BATCH_SIZE + 1}...")
        
        batch_outputs = generate_summary_batch(
            batch_chunks,
            max_new_tokens=200,
            is_final=False
        )
        
        for part_sum in batch_outputs:
            if "PASS" not in part_sum and len(part_sum) > 10:
                partials.append(part_sum)

    # Nối các mảnh lại bằng dấu xuống dòng
    merged = "\n".join(partials)

    # PASS 1: Đẩy vào vòng chung kết để lấy danh sách gạch đầu dòng sạch rác
    flat_summary = generate_summary(merged, max_new_tokens=1000, is_final=True)

    # PASS 2: Gom nhóm và chuyển đổi thành đoạn văn xuôi
    final_grouped_paragraphs = dynamic_clustering_to_paragraphs(flat_summary)

    return final_grouped_paragraphs


def dynamic_clustering_to_paragraphs(flat_bullets):
    """
    BƯỚC 2 (POST-PROCESSING): Nhận vào danh sách gạch đầu dòng phẳng,
    tự động nhận diện chuyên mục và hành văn lại thành các đoạn văn xuôi liên kết.
    """
    system_prompt = """Bạn là Chuyên gia soạn thảo văn bản hành chính.
Nhiệm vụ: Hãy đọc danh sách các sự kiện dưới đây và thực hiện:
1. Tự nhận diện 2 đến 4 chủ đề/lĩnh vực chính bao quát nội dung.
2. Gom các sự kiện liên quan vào đúng chủ đề.
3. HÀNH VĂN THÀNH ĐOẠN (PARAGRAPH): Viết lại các ý trong mỗi chủ đề thành một đoạn văn xuôi hoàn chỉnh. Nối các sự kiện bằng dấu chấm và từ nối logic (tuy nhiên, do, sau đó...).
4. TUYỆT ĐỐI KHÔNG dùng gạch đầu dòng. KHÔNG tự bịa thêm thông tin ngoài danh sách.

Định dạng đầu ra BẮT BUỘC:
[TÊN CHỦ ĐỀ TỰ NHẬN DIỆN]:
(Nội dung đoạn văn xuôi viết liền mạch, nối với nhau bằng dấu câu hợp lý, không xuống dòng giữa chừng)

[TÊN CHỦ ĐỀ KHÁC]:
(Nội dung đoạn văn xuôi...)
"""
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Danh sách sự kiện:\n{flat_bullets}\n\nHãy gom nhóm và viết thành các đoạn văn:"}
    ]

    text_input = _qwen_tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)

    inputs = _qwen_tokenizer(text_input, return_tensors="pt", truncation=True, max_length=2500).to(_qwen_model.device)

    with torch.no_grad():
        outputs = _qwen_model.generate(
            **inputs,
            max_new_tokens=1000,
            repetition_penalty=1.1,
            do_sample=False, # Khóa Greedy Decoding để chống ảo giác
            pad_token_id=_qwen_tokenizer.eos_token_id,
            eos_token_id=_qwen_tokenizer.eos_token_id
        )

    generated_ids = [output_ids[len(input_ids):] for input_ids, output_ids in zip(inputs.input_ids, outputs)]
    final_text = _qwen_tokenizer.batch_decode(generated_ids, skip_special_tokens=True)[0]

    return final_text.replace("<|im_end|>", "").strip()



# ========================= GUI (giữ nguyên hoàn toàn) =========================
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

class MeetingApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Hệ thống Tóm tắt Biên bản Cuộc họp")
        self.geometry("1100x800")
        self.minsize(850, 600)
        self._center_window()

        self.audio_file      = None
        self.audio_timestamp = None
        self.recording       = False
        self.audio_data      = []
        self.sample_rate     = 16000
        self.record_thread   = None
        self.raw_asr_text    = ""

        self._build_ui()
        # self._start_model_loading()

    def _center_window(self):
        w, h = 1000, 700
        x = (self.winfo_screenwidth()  - w) // 2
        y = (self.winfo_screenheight() - h) // 2
        self.geometry(f"{w}x{h}+{x}+{y}")

    def _build_ui(self):
        self.configure(fg_color="#F8FAFC")

        # Header
        header = ctk.CTkFrame(self, fg_color="transparent")
        header.pack(pady=(10, 5), padx=40, fill="x")
        ctk.CTkLabel(header, text="HỆ THỐNG TÓM TẮT BIÊN BẢN",
                     font=ctk.CTkFont(family="Segoe UI", size=26, weight="bold"), text_color="#000000").pack(anchor="w")
        self.model_status = ctk.CTkLabel(
            header, text="Trạng thái hệ thống: Đang khởi tạo mô hình...",
            font=ctk.CTkFont(size=18, weight="bold"), text_color="#000000")
        self.model_status.pack(anchor="w")

        # Loading banner
        self.model_loading_frame = ctk.CTkFrame(
            self, fg_color="#F0F8FF", border_width=1, border_color="#BFDBFE", corner_radius=8)
        self.model_loading_frame.pack(pady=0, padx=40, fill="x")
        
        row_frame = ctk.CTkFrame(self.model_loading_frame, fg_color="transparent")
        row_frame.pack(fill="x", padx=20, pady=2)
        
        self.model_loading_label = ctk.CTkLabel(
            row_frame, text="Đang load mô hình...",
            font=ctk.CTkFont(size=18, weight="bold"), text_color="#000000")
        self.model_loading_label.pack(side="left")
        
        self.model_progress = ctk.CTkProgressBar(self.model_loading_frame, height=6, mode="indeterminate", progress_color="#1E3A8A")
        self.model_progress.pack(fill="x", padx=20, pady=(0, 6))
        self.model_progress.start()

        # Input
        in_frame = ctk.CTkFrame(self, fg_color="#FFFFFF", border_width=1, border_color="#E2E8F0", corner_radius=8)
        in_frame.pack(pady=4, padx=40, fill="x")
        ctk.CTkLabel(in_frame, text="NGUỒN DỮ LIỆU ÂM THANH",
                     font=ctk.CTkFont(size=18, weight="bold"), text_color="#000000").pack(anchor="w", padx=24, pady=(8, 4))

        btn_row = ctk.CTkFrame(in_frame, fg_color="transparent")
        btn_row.pack(padx=20, pady=2, fill="x")
        self.btn_upload = ctk.CTkButton(
            btn_row, text="Tải tệp âm thanh", height=42,
            font=ctk.CTkFont(size=18, weight="bold"), command=self._upload_file)
        self.btn_upload.pack(side="left", padx=4, expand=True, fill="x")
        self.btn_record = ctk.CTkButton(
            btn_row, text="Ghi âm trực tiếp", height=42,
            font=ctk.CTkFont(size=18, weight="bold"), command=self._toggle_recording)
        self.btn_record.pack(side="left", padx=4, expand=True, fill="x")

        self.file_label = ctk.CTkLabel(
            in_frame, text="Chưa có tệp nào được chọn.",
            font=ctk.CTkFont(size=18, weight="bold"), text_color="#000000")
        self.file_label.pack(anchor="w", padx=24, pady=(2, 0))
        
        self.suggestion_label = ctk.CTkLabel(
            in_frame, text="", 
            font=ctk.CTkFont(size=16, weight="bold"), text_color="#D97706")
        self.suggestion_label.pack(anchor="w", padx=24, pady=(0, 6))

        # Method
        method_frame = ctk.CTkFrame(self, fg_color="#FFFFFF", border_width=1, border_color="#E2E8F0", corner_radius=8)
        method_frame.pack(pady=4, padx=40, fill="x")
        ctk.CTkLabel(method_frame, text="CHỌN CÁCH TÓM TẮT",
                     font=ctk.CTkFont(size=18, weight="bold"), text_color="#000000").pack(anchor="w", padx=24, pady=(8, 2))

        self.method_var = ctk.StringVar(value="textrank")
        row = ctk.CTkFrame(method_frame, fg_color="transparent")
        row.pack(padx=20, pady=(0, 8), fill="x")
        ctk.CTkRadioButton(row, text="Cuộc họp ngắn", variable=self.method_var, value="textrank", font=ctk.CTkFont(size=18, weight="bold"), text_color="#000000").pack(side="left", padx=8)
        ctk.CTkRadioButton(row, text="Cuộc họp dài", variable=self.method_var, value="chunking", font=ctk.CTkFont(size=18, weight="bold"), text_color="#000000").pack(side="left", padx=24)

        self.progress = ctk.CTkProgressBar(self, height=16, progress_color="#D97706") # Đậm màu hơn
        self.progress.pack(pady=(8, 4), padx=40, fill="x")
        self.process_btn = ctk.CTkButton(
            self, text="BẮT ĐẦU XỬ LÝ", height=48,
            font=ctk.CTkFont(size=20, weight="bold"),
            command=self._start_processing, state="disabled")
        self.process_btn.pack(pady=2, padx=40, fill="x")

        # --- Đưa Status & Save Buttons xuống dưới đáy (Bottom) để không bao giờ bị lấp ---
        save_row = ctk.CTkFrame(self, fg_color="transparent")
        save_row.pack(side="bottom", pady=(4, 15), padx=40, fill="x")
        self.save_word_btn = ctk.CTkButton(save_row, text="Xuất Biên bản Word", width=180, height=36, font=ctk.CTkFont(size=16, weight="bold"), state="disabled", command=self._save_word)
        self.save_word_btn.pack(side="left")

        self.user_status = ctk.CTkLabel(self, text="", font=ctk.CTkFont(size=18, weight="bold"), text_color="#000000")
        self.user_status.pack(side="bottom", pady=(0, 5), padx=40, anchor="w")

        # Tabs: Tabview được Pack cuối cùng với expand=True để nó tự điền vào khoảng trống ở giữa
        self.tabview = ctk.CTkTabview(self, height=200)
        self.tabview.pack(side="top", pady=6, padx=40, fill="both", expand=True)
        try:
            self.tabview._segmented_button.configure(font=ctk.CTkFont(size=16, weight="bold"))
        except:
            pass
        self.tabview.add("Biên bản")
        self.tabview.add("Nhật ký (Log)")

        self.result_text = ctk.CTkTextbox(self.tabview.tab("Biên bản"), font=ctk.CTkFont(size=22, weight="bold"))
        self.result_text.pack(fill="both", expand=True, padx=2, pady=2)

        self.log_text = ctk.CTkTextbox(self.tabview.tab("Nhật ký (Log)"), font=ctk.CTkFont(family="Consolas", size=16, weight="bold"), fg_color="#FFFFFF", text_color="#000000")
        self.log_text.pack(fill="both", expand=True, padx=2, pady=2)

        import sys
        class RedirectText:
            def __init__(self, app, text_widget):
                self.app = app
                self.output = text_widget
            def write(self, string):
                self.app.after(0, self._write, string)
            def _write(self, string):
                self.output.insert("end", string)
                self.output.see("end")
            def flush(self):
                pass
        sys.stdout = RedirectText(self, self.log_text)

    # ─── Model loading ───────────────────────────────────────────────────────
    def _start_model_loading(self):
        def _run():
            ok = load_models()
            self.after(0, self._on_models_loaded, ok)

        threading.Thread(target=_run, daemon=True).start()

    def _on_models_loaded(self, ok):
        self.model_progress.stop()
        if ok:
            self.model_loading_frame.configure(fg_color="#F0FDF4", border_color="#86EFAC")
            self.model_loading_label.configure(text="Đã load xong!", text_color="#15803D")
            self.model_progress.configure(progress_color="#22C55E")
            self.model_progress.set(1.0)
            self.model_status.configure(text="Trạng thái hệ thống: Sẵn sàng", text_color="#059669")
            self.process_btn.configure(state="normal")
            self.after(3000, self._hide_loading_banner)
        else:
            self.model_loading_frame.configure(fg_color="#FEF2F2", border_color="#FECACA")
            self.model_loading_label.configure(text="❌ Nạp mô hình thất bại", text_color="#DC2626")

    def _hide_loading_banner(self):
        self.model_loading_frame.pack_forget()

    # ─── Audio ───────────────────────────────────────────────────────────────
    def _upload_file(self):
        path = filedialog.askopenfilename(filetypes=[("Audio files", "*.wav *.mp3 *.m4a *.ogg *.flac")])
        if path:
            self._clear_result()
            self.audio_file      = path
            self.audio_timestamp = datetime.fromtimestamp(os.path.getctime(path))
            
            try:
                import librosa
                try:
                    duration = librosa.get_duration(path=path)
                except:
                    duration = librosa.get_duration(filename=path)
                mins = int(duration // 60)
                secs = int(duration % 60)
                time_str = f"{mins}p {secs}s"
                sugg = "cuộc họp ngắn" if duration < 600 else "cuộc họp dài"
            except:
                time_str = "Không rõ"
                sugg = ""

            self.file_label.configure(text=f"Tệp đã chọn: {os.path.basename(path)} | Thời gian: {time_str}", text_color="#0F172A")
            self.user_status.configure(text=f"Đã tải file: {os.path.basename(path)}")
            if sugg:
                self.suggestion_label.configure(text=f"Gợi ý: Chọn tóm tắt {sugg}", text_color="#D97706")
            else:
                self.suggestion_label.configure(text="")

    def _update_record_time(self, start_time):
        if self.recording:
            elapsed = int(time.time() - start_time)
            mins = elapsed // 60
            secs = elapsed % 60
            self.file_label.configure(text=f"Đang ghi âm... {mins:02d}:{secs:02d}", text_color="#2563EB")
            self.after(1000, self._update_record_time, start_time)

    def _toggle_recording(self):
        if not self.recording:
            self.recording = True
            self._clear_result()
            self.btn_record.configure(text="Dừng ghi âm", fg_color="#FEE2E2", text_color="#DC2626")
            self.file_label.configure(text="Đang ghi âm... 00:00", text_color="#2563EB")
            self.suggestion_label.configure(text="")
            self.audio_data = []
            self.record_thread = threading.Thread(target=self._record_audio, daemon=True)
            self.record_thread.start()
            self._update_record_time(time.time())
        else:
            self.recording = False
            self.btn_record.configure(text="Ghi âm trực tiếp", fg_color="#F1F5F9", text_color="#0F172A")

    def _record_audio(self):
        import sounddevice as sd
        import soundfile as sf
        try:
            with sd.InputStream(samplerate=self.sample_rate, channels=1, dtype="float32") as stream:
                while self.recording:
                    data, _ = stream.read(1024)
                    self.audio_data.append(data)
        finally:
            if self.audio_data:
                audio_arr = np.concatenate(self.audio_data, axis=0)
                ts = datetime.now()
                self.audio_file = f"recording_{ts.strftime('%Y%m%d_%H%M%S')}.wav"
                sf.write(self.audio_file, audio_arr, self.sample_rate)
                self.audio_timestamp = ts

                duration = len(audio_arr) / self.sample_rate
                mins = int(duration // 60)
                secs = int(duration % 60)
                time_str = f"{mins}p {secs}s"
                sugg = "ngắn" if duration < 600 else "dài"

                self.after(0, lambda f=self.audio_file, t=time_str: self.file_label.configure(
                    text=f"Đã ghi âm: {f} | Thời gian: {t}", text_color="#0F172A"))
                # self.after(0, lambda f=self.audio_file: self.user_status.configure(
                #     text=f"Đã ghi âm xong: {os.path.basename(f)}"))
                self.after(0, lambda s=sugg: self.suggestion_label.configure(
                    text=f"Gợi ý: Chọn tóm tắt cuộc họp {s}", text_color="#D97706"))

    # ─── Processing ──────────────────────────────────────────────────────────
    def _start_processing(self):
        if not self.audio_file:
            msgbox.showwarning("Cảnh báo", "Vui lòng chọn hoặc ghi âm file trước!")
            return
        if not _models_loaded:
            msgbox.showwarning("Cảnh báo", "Hệ thống đang khởi tạo...")
            return

        method = "TextRank" if self.method_var.get() == "textrank" else "Map-Reduce Chunking"
        self.user_status.configure(
            text=f"Đang xử lý • Phương pháp: {method} • File: {os.path.basename(self.audio_file)}")
        self.process_btn.configure(text="ĐANG XỬ LÝ...", state="disabled")
        self.progress.set(0)
        self._clear_result()
        threading.Thread(target=self._pipeline_thread, daemon=True).start()

    def _pipeline_thread(self):
        try:
            def progress_cb(done, total):
                self.after(0, lambda: self.progress.set(0.1 + (done / total) * 0.4))

            raw_text = transcribe_audio_pipeline(self.audio_file, progress_cb)
            self.raw_asr_text = raw_text
            self.after(0, lambda: self.progress.set(0.55))

            print("[Giai đoạn 2/3] Đang làm sạch văn bản (Qwen)...")
            clean_text = restore_punctuation_and_clean(raw_text)
            self.after(0, lambda: self.progress.set(0.75))

            method = self.method_var.get()
            method_name = "TextRank" if method == "textrank" else "Map-Reduce Chunking"
            print(f"[Giai đoạn 3/3] Đang tóm tắt ({method_name})...")
            summary = enhanced_textrank(clean_text) if method == "textrank" else chunking_method(clean_text)

            print("✔ Hoàn tất xử lý!")
            self.after(0, lambda: self.progress.set(1.0))
            self.after(0, self._finish_processing, summary)

        except Exception as e:
            tb = traceback.format_exc()
            print(f"[LỖI XỬ LÝ]\n{tb}")
            self.after(0, self._on_error, str(e))

    def _finish_processing(self, summary):
        self.result_text.delete("0.0", "end")
        self.result_text.insert("0.0", summary)
        self.process_btn.configure(text="BẮT ĐẦU XỬ LÝ", state="normal")
        self.save_word_btn.configure(state="normal")
        self.user_status.configure(text="Hoàn tất! Bạn có thể xem kết quả hoặc xuất Biên bản Word.")

    def _on_error(self, msg=""):
        self.process_btn.configure(text="BẮT ĐẦU XỬ LÝ", state="normal")
        self.tabview.set("Nhật ký (Log)")
        msgbox.showerror("Lỗi xử lý", f"Có lỗi xảy ra:\n{msg}\n\nXem chi tiết ở tab Nhật ký (Log).")

    def _save_word(self):
        content = self.result_text.get("0.0", "end").strip()
        if not content:
            msgbox.showwarning("Cảnh báo", "Chưa có nội dung biên bản để xuất!")
            return
            
        # Gọi thẳng hàm định dạng Word từ văn bản có sẵn trên GUI
        doc = tao_form_bien_ban_mau(content)
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"BienBan_ChinhThuc_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            filetypes=[("Word documents", "*.docx")]
        )
        if path:
            doc.save(path)
            msgbox.showinfo("Thành công",
                            f"Đã xuất biên bản Word mẫu:\n{path}\n\nBạn có thể mở file để chỉnh sửa thêm.")

    def _clear_result(self):
        self.result_text.delete("0.0", "end")
        self.save_word_btn.configure(state="disabled")
        self.progress.set(0)
        self.user_status.configure(text="")


if __name__ == "__main__":
    app = MeetingApp()
    app.mainloop()
